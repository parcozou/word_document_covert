"""Generate and audit a DOCX from the sample Coze payload."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn

from docx_converter import convert_markdown_to_docx, safe_file_stem


DEFAULT_SAMPLE_PAYLOAD = {
    "formatted_markdown": (
        "# DOCX Conversion Check\n\n"
        "## 1. Executive Summary\n\n"
        "### Evidence Figure\n\n"
        "![Validation Figure](https://lf-bot-studio-plugin-resource.coze.cn/invalid-test-chart.png)\n"
        "*Note: This note should be rendered beneath the figure.*\n"
        "Analysis follows the figure note.\n\n"
        "Evidence is supported by the first source [1].\n\n"
        "- This bullet point should be justified when it wraps across lines in the Word report.\n\n"
        "| Metric | Value |\n"
        "| --- | ---: |\n"
        "| Revenue | RMB 100.0m |\n"
        "| Margin | 25.0% |\n\n"
        "## 2. Conclusion\n\n"
        "The standalone validation payload confirms native table conversion.\n\n"
        "## References\n"
        "[1] First test source.\n"
        "[2] Second test source."
    ),
    "title": "DOCX Conversion Check",
}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--input",
        type=Path,
        help="Optional JSON payload file with formatted_markdown and title fields.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path(__file__).resolve().parent / "generated_files",
    )
    parser.add_argument("--skip-images", action="store_true")
    args = parser.parse_args()

    payload = (
        json.loads(args.input.read_text(encoding="utf-8"))
        if args.input
        else DEFAULT_SAMPLE_PAYLOAD
    )
    output = args.output_dir / f"{safe_file_stem(payload['title'])}_sample.docx"
    result = convert_markdown_to_docx(
        payload["formatted_markdown"],
        payload["title"],
        output,
        include_images=not args.skip_images,
    )

    document = Document(str(output))
    invalid_table_cells = [
        cell.text.strip()[:80]
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip().startswith(("#", "---", "!["))
    ]
    suspicious_headings = [
        paragraph.text.strip()[:100]
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
        and len(paragraph.text.split()) > 18
    ]
    embedded_list_lines = [
        paragraph.text.strip()[:100]
        for paragraph in document.paragraphs
        if "\n-" in paragraph.text
        or re.search(r"\n\d+[.)]\s", paragraph.text)
    ]
    document_xml = document._element.xml
    normal_style = document.styles["Normal"]
    has_contents_field = 'TOC \\o "1-3"' in document_xml
    first_body_heading = next(
        (
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.style.name.startswith("Heading") and paragraph.text.strip()
        ),
        "",
    )
    contents_are_populated = (
        bool(first_body_heading)
        and sum(
            paragraph.text.strip() == first_body_heading
            for paragraph in document.paragraphs
        )
        >= 2
    )
    reference_items = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if re.match(r"^\[\d+\]\s+", paragraph.text.strip())
    ]
    heading_font_errors = [
        paragraph.text.strip()[:100]
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
        and any(run.text and run.font.name != "Times New Roman" for run in paragraph.runs)
    ]
    table_format_errors = [
        cell.text.strip()[:80]
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if paragraph.paragraph_format.space_before != 0
        or paragraph.paragraph_format.space_after != 0
        or paragraph.paragraph_format.line_spacing != 1.0
        or any(run.text and run.font.size != 12 * 12700 for run in paragraph.runs)
    ]
    table_rule_errors = []
    for table in document.tables:
        last_row_index = len(table.rows) - 1
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                borders = cell._tc.tcPr.find(qn("w:tcBorders"))
                top = borders.find(qn("w:top")).get(qn("w:val"))
                bottom = borders.find(qn("w:bottom")).get(qn("w:val"))
                top_size = borders.find(qn("w:top")).get(qn("w:sz"))
                bottom_size = borders.find(qn("w:bottom")).get(qn("w:sz"))
                expected_top = "single" if row_index == 0 else "nil"
                expected_bottom = (
                    "single" if row_index in {0, last_row_index} else "nil"
                )
                expected_top_size = "12" if row_index == 0 else "0"
                expected_bottom_size = (
                    "12" if row_index == last_row_index else "4" if row_index == 0 else "0"
                )
                if (
                    top != expected_top
                    or bottom != expected_bottom
                    or top_size != expected_top_size
                    or bottom_size != expected_bottom_size
                ):
                    table_rule_errors.append(cell.text.strip()[:80])
    normal_auto_spacing = all(
        attribute in normal_style._element.xml
        for attribute in ('w:beforeAutospacing="1"', 'w:afterAutospacing="1"')
    )
    body_title_present = any(
        paragraph.style.name.startswith("Heading")
        and paragraph.text.strip() == payload["formatted_markdown"].splitlines()[0].lstrip("#").strip()
        for paragraph in document.paragraphs
    )
    bullet_alignment_errors = [
        paragraph.text.strip()[:80]
        for paragraph in document.paragraphs
        if paragraph.style.name in {"List Bullet", "List Number", "List Paragraph"}
        and paragraph.alignment != 3
    ]
    cover_is_vertically_centered = (
        '<w:vAlign w:val="center"/>' in document.sections[0]._sectPr.xml
    )
    markdown_image_titles = {
        match.strip()
        for match in re.findall(r"!\[([^\]]+)\]\([^)]+\)", payload["formatted_markdown"])
        if match.strip()
    }
    duplicate_chart_title_paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if result.image_count
        and paragraph.text.strip() in markdown_image_titles
        and not paragraph.text.strip().startswith("[Figure unavailable:")
    ]
    citation_links_present = 'w:anchor="reference_1"' in document_xml
    bookmark_present = 'w:name="reference_1"' in document_xml
    superscript_citation_present = (
        "Evidence is supported by the first source" not in payload["formatted_markdown"]
        or "<w:vertAlign w:val=\"superscript\"/>" in document_xml
    )
    heading_hierarchy_errors = []
    heading_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip()
    ]
    for paragraph in heading_paragraphs:
        text = paragraph.text.strip()
        if text == payload["formatted_markdown"].splitlines()[0].lstrip("#").strip():
            expected_style = "Heading 1"
        elif re.match(r"^\d+\.\d+(?:\.\d+)?\s", text):
            expected_style = "Heading 3"
        elif re.match(r"^\d+\.\s", text) or text.lower() in {
            "bottom line",
            "references",
            "references / sources",
            "disclaimer",
        }:
            expected_style = "Heading 2"
        else:
            expected_style = None
        if expected_style and paragraph.style.name != expected_style:
            heading_hierarchy_errors.append(
                f"{text[:80]} ({paragraph.style.name}; expected {expected_style})"
            )
        if paragraph.style.name == "Heading 3" and not re.match(
            r"^\d+\.\d+(?:\.\d+)?\s", text
        ):
            heading_hierarchy_errors.append(
                f"{text[:80]} (unnumbered Heading 3)"
            )
    page_break_errors = []
    report_seen = False
    waiting_for_first_major_section = False
    for paragraph in heading_paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 1":
            expected_break = report_seen
            if bool(paragraph.paragraph_format.page_break_before) != expected_break:
                page_break_errors.append(text[:80])
            report_seen = True
            waiting_for_first_major_section = True
        elif paragraph.style.name == "Heading 2":
            expected_break = not waiting_for_first_major_section
            if bool(paragraph.paragraph_format.page_break_before) != expected_break:
                page_break_errors.append(text[:80])
            waiting_for_first_major_section = False
    figure_note_order_ok = True
    if not args.input:
        paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs]
        expected_sequence = [
            "[Figure unavailable: Validation Figure]",
            "Note: This note should be rendered beneath the figure.",
            "Analysis follows the figure note.",
        ]
        locations = [
            paragraph_text.index(value) if value in paragraph_text else -1
            for value in expected_sequence
        ]
        figure_note_order_ok = (
            -1 not in locations and locations == sorted(locations)
        )
    has_excluded_academic_forms = any(
        phrase in "\n".join(paragraph.text for paragraph in document.paragraphs)
        for phrase in ("Contribution Statement", "Declaration")
    )
    audit = {
        "output_path": str(output.resolve()),
        "native_word_tables": len(document.tables),
        "reported_tables": result.table_count,
        "embedded_images": result.image_count,
        "skipped_images": result.skipped_image_count,
        "headings": result.heading_count,
        "invalid_table_cells": invalid_table_cells,
        "suspicious_headings": suspicious_headings,
        "embedded_list_lines": embedded_list_lines,
        "sections": len(document.sections),
        "normal_font": normal_style.font.name,
        "line_spacing": normal_style.paragraph_format.line_spacing,
        "has_contents_field": has_contents_field,
        "contents_are_populated": contents_are_populated,
        "reference_item_paragraphs": len(reference_items),
        "heading_font_errors": heading_font_errors,
        "table_format_errors": table_format_errors,
        "table_rule_errors": table_rule_errors,
        "normal_auto_spacing": normal_auto_spacing,
        "body_title_present": body_title_present,
        "bullet_alignment_errors": bullet_alignment_errors,
        "cover_is_vertically_centered": cover_is_vertically_centered,
        "duplicate_chart_title_paragraphs": duplicate_chart_title_paragraphs,
        "citation_links_present": citation_links_present,
        "bookmark_present": bookmark_present,
        "superscript_citation_present": superscript_citation_present,
        "heading_hierarchy_errors": heading_hierarchy_errors,
        "page_break_errors": page_break_errors,
        "figure_note_order_ok": figure_note_order_ok,
    }
    if not document.tables or len(document.tables) != result.table_count:
        raise RuntimeError(f"Table audit failed: {audit}")
    if invalid_table_cells:
        raise RuntimeError(f"Content was incorrectly embedded in a table: {audit}")
    if suspicious_headings:
        raise RuntimeError(f"Long narrative content was incorrectly styled as a heading: {audit}")
    if embedded_list_lines:
        raise RuntimeError(f"List items remained embedded inside body paragraphs: {audit}")
    if len(document.sections) != 3 or not has_contents_field:
        raise RuntimeError(f"Front matter audit failed: {audit}")
    if not contents_are_populated:
        raise RuntimeError(f"Contents entries were not generated before Word field refresh: {audit}")
    if normal_style.font.name != "Times New Roman" or normal_style.paragraph_format.line_spacing != 1.5:
        raise RuntimeError(f"Typography audit failed: {audit}")
    if heading_font_errors:
        raise RuntimeError(f"Heading font audit failed: {audit}")
    if table_format_errors:
        raise RuntimeError(f"Table typography or spacing audit failed: {audit}")
    if table_rule_errors:
        raise RuntimeError(f"Table rule-border audit failed: {audit}")
    if not normal_auto_spacing:
        raise RuntimeError(f"Body automatic-spacing audit failed: {audit}")
    if not body_title_present:
        raise RuntimeError(f"Report section title was removed from body output: {audit}")
    if bullet_alignment_errors:
        raise RuntimeError(f"Bullet justification audit failed: {audit}")
    if not cover_is_vertically_centered:
        raise RuntimeError(f"Cover vertical-centering audit failed: {audit}")
    if duplicate_chart_title_paragraphs:
        raise RuntimeError(f"Duplicate chart-title paragraph audit failed: {audit}")
    if heading_hierarchy_errors:
        raise RuntimeError(f"Heading hierarchy audit failed: {audit}")
    if page_break_errors:
        raise RuntimeError(f"Major-section page-break audit failed: {audit}")
    if not figure_note_order_ok:
        raise RuntimeError(f"Figure-note order audit failed: {audit}")
    if not args.input and len(reference_items) != 2:
        raise RuntimeError(f"Reference list paragraph audit failed: {audit}")
    if not args.input and not all(
        (citation_links_present, bookmark_present, superscript_citation_present)
    ):
        raise RuntimeError(f"Citation cross-reference audit failed: {audit}")
    if has_excluded_academic_forms:
        raise RuntimeError(f"Excluded academic forms were generated: {audit}")
    for section in document.sections:
        if round(section.page_width.inches, 2) != 8.27 or round(section.page_height.inches, 2) != 11.69:
            raise RuntimeError(f"Page layout audit failed: {audit}")
    print(json.dumps(audit, indent=2))


if __name__ == "__main__":
    main()
