"""Convert a Markdown report into a formatted DOCX with native Word tables."""

from __future__ import annotations

import gc
from dataclasses import dataclass
from datetime import date
from pathlib import Path
import os
import re
import tempfile
from urllib.parse import urlparse

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor
import markdown
from PIL import Image, UnidentifiedImageError
import requests


REPORT_FONT = "Times New Roman"
TABLE_WIDTH_IN = 5.75
TABLE_WIDTH_DXA = 8280
TABLE_INDENT_DXA = 0
DEFAULT_IMAGE_HOSTS = "lf-bot-studio-plugin-resource.coze.cn"
DEFAULT_IMAGE_MAX_WIDTH_PX = 1400
DEFAULT_IMAGE_MAX_HEIGHT_PX = 1000
DEFAULT_IMAGE_DOWNLOAD_MAX_BYTES = 12 * 1024 * 1024
TABLE_DELIMITER_CELL = re.compile(r"^:?-{3,}:?$")
LIST_ITEM_LINE = re.compile(r"^\s*(?:[-+*]\s+|\d+[.)]\s+)")
REFERENCE_HEADING_LINE = re.compile(
    r"^\s*#{1,6}\s*(?:\d+[.)]?\s*)?(?:references?|reference\s+list|sources?)(?:\s*/\s*sources?)?\s*$",
    re.IGNORECASE,
)
REFERENCE_ITEM_LINE = re.compile(r"^\s*\[\d+\]\s+")
ANY_HEADING_LINE = re.compile(r"^\s*#{1,6}\s+")
CITATION_MARKER = re.compile(r"\[(\d+)\]")
MARKDOWN_IMAGE_LINE = re.compile(r"^\s*!\[[^\]]*\]\([^)]+\)\s*$")
NOTE_LINE = re.compile(r"^\s*[*_]*\s*Note\s*:", re.IGNORECASE)
MARKDOWN_HEADING = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.+?)\s*$")
MAJOR_HEADING_TEXT = re.compile(r"^(?P<number>\d+)\.\s+")
SUBSECTION_HEADING_TEXT = re.compile(r"^(?P<major>\d+)\.(?P<minor>\d+)(?:\.\d+)?\s+")


@dataclass
class ConversionResult:
    output_path: Path
    table_count: int = 0
    image_count: int = 0
    skipped_image_count: int = 0
    heading_count: int = 0


@dataclass
class RenderState:
    report_count: int = 0
    waiting_for_first_major_section: bool = False


def _table_cells(line: str) -> list[str] | None:
    stripped = line.strip()
    if "|" not in stripped:
        return None
    cells = stripped.strip("|").split("|")
    return [cell.strip() for cell in cells]


def _is_table_delimiter(line: str) -> bool:
    cells = _table_cells(line)
    return bool(
        cells
        and len(cells) >= 2
        and all(TABLE_DELIMITER_CELL.fullmatch(cell) for cell in cells)
    )


def _normalise_table_boundaries(markdown_text: str) -> str:
    """Close LLM-written pipe tables even when a following blank line is missing."""
    lines = markdown_text.splitlines()
    normalised: list[str] = []
    index = 0
    while index < len(lines):
        starts_table = (
            index + 1 < len(lines)
            and _table_cells(lines[index]) is not None
            and _is_table_delimiter(lines[index + 1])
        )
        if not starts_table:
            normalised.append(lines[index])
            index += 1
            continue

        if normalised and normalised[-1].strip():
            normalised.append("")

        expected_columns = len(_table_cells(lines[index + 1]) or [])
        normalised.extend([lines[index], lines[index + 1]])
        index += 2
        while index < len(lines):
            cells = _table_cells(lines[index])
            if cells is None or len(cells) != expected_columns:
                break
            normalised.append(lines[index])
            index += 1

        if normalised[-1].strip():
            normalised.append("")

    return "\n".join(normalised)


def _normalise_horizontal_rules(markdown_text: str) -> str:
    """Prevent section dividers from converting the prior paragraph to an H2."""
    lines = markdown_text.splitlines()
    normalised: list[str] = []
    for index, line in enumerate(lines):
        if line.strip() == "---":
            if normalised and normalised[-1].strip():
                normalised.append("")
            normalised.append(line)
            if index + 1 < len(lines) and lines[index + 1].strip():
                normalised.append("")
            continue
        normalised.append(line)
    return "\n".join(normalised)


def _normalise_list_boundaries(markdown_text: str) -> str:
    """Give compact LLM-authored list lines the paragraph boundaries Markdown expects."""
    lines = markdown_text.splitlines()
    normalised: list[str] = []
    previous_was_item = False
    for line in lines:
        is_item = LIST_ITEM_LINE.match(line) is not None
        if is_item and normalised and normalised[-1].strip() and not previous_was_item:
            normalised.append("")
        if not is_item and previous_was_item and line.strip():
            normalised.append("")
        normalised.append(line)
        previous_was_item = is_item
    return "\n".join(normalised)


def _normalise_reference_boundaries(markdown_text: str) -> str:
    """Keep individually numbered source entries as separate Word paragraphs."""
    lines = markdown_text.splitlines()
    normalised: list[str] = []
    in_references = False
    for line in lines:
        if REFERENCE_HEADING_LINE.match(line):
            in_references = True
        elif in_references and ANY_HEADING_LINE.match(line):
            in_references = False
        if in_references and REFERENCE_ITEM_LINE.match(line):
            if normalised and normalised[-1].strip():
                normalised.append("")
        normalised.append(line)
    return "\n".join(normalised)


def _normalise_figure_boundaries(markdown_text: str) -> str:
    """Keep standalone chart images and immediately following notes in source order."""
    lines = markdown_text.splitlines()
    normalised: list[str] = []
    for index, line in enumerate(lines):
        needs_own_paragraph = bool(
            MARKDOWN_IMAGE_LINE.match(line) or NOTE_LINE.match(line)
        )
        if needs_own_paragraph and normalised and normalised[-1].strip():
            normalised.append("")
        normalised.append(line)
        if (
            needs_own_paragraph
            and index + 1 < len(lines)
            and lines[index + 1].strip()
        ):
            normalised.append("")
    return "\n".join(normalised)


def _normalise_subsection_numbers(markdown_text: str) -> str:
    """Number missing subsection labels below numbered major report sections."""
    lines = markdown_text.splitlines()
    normalised: list[str] = []
    major_number: int | None = None
    major_level: int | None = None
    subsection_number = 0
    for line in lines:
        match = MARKDOWN_HEADING.match(line)
        if not match:
            normalised.append(line)
            continue

        level = len(match.group("hashes"))
        text = match.group("text").strip()
        if level == 1:
            major_number = None
            major_level = None
            subsection_number = 0

        major_match = MAJOR_HEADING_TEXT.match(text)
        if major_match and not SUBSECTION_HEADING_TEXT.match(text):
            major_number = int(major_match.group("number"))
            major_level = level
            subsection_number = 0
            normalised.append(line)
            continue

        subsection_match = SUBSECTION_HEADING_TEXT.match(text)
        if subsection_match:
            if major_number == int(subsection_match.group("major")):
                subsection_number = max(
                    subsection_number, int(subsection_match.group("minor"))
                )
            normalised.append(line)
            continue

        if major_number is not None and major_level is not None and level > major_level:
            subsection_number += 1
            line = (
                f"{match.group('hashes')} {major_number}.{subsection_number} {text}"
            )
        normalised.append(line)
    return "\n".join(normalised)


def _normalise_report_markdown(markdown_text: str) -> str:
    with_numbered_subsections = _normalise_subsection_numbers(markdown_text)
    with_closed_tables = _normalise_table_boundaries(with_numbered_subsections)
    with_list_boundaries = _normalise_list_boundaries(with_closed_tables)
    with_reference_boundaries = _normalise_reference_boundaries(with_list_boundaries)
    with_figure_boundaries = _normalise_figure_boundaries(with_reference_boundaries)
    return _normalise_horizontal_rules(with_figure_boundaries)


def _safe_hosts() -> set[str]:
    configured = os.getenv("IMAGE_HOST_ALLOWLIST", DEFAULT_IMAGE_HOSTS)
    return {host.strip().lower() for host in configured.split(",") if host.strip()}


def _positive_int_from_env(name: str, default: int) -> int:
    try:
        value = int(os.getenv(name, str(default)))
    except ValueError:
        return default
    return value if value > 0 else default


def _font_name(element, name: str = REPORT_FONT) -> None:
    element.font.name = name
    properties = element._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:cs"), name)
    for theme_attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        qualified = qn(f"w:{theme_attribute}")
        if qualified in fonts.attrib:
            del fonts.attrib[qualified]


def _set_auto_spacing(element, enabled: bool) -> None:
    paragraph_properties = element._element.get_or_add_pPr()
    spacing = paragraph_properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        paragraph_properties.append(spacing)
    value = "1" if enabled else "0"
    spacing.set(qn("w:beforeAutospacing"), value)
    spacing.set(qn("w:afterAutospacing"), value)
    if enabled:
        for manual_spacing in ("before", "after"):
            qualified = qn(f"w:{manual_spacing}")
            if qualified in spacing.attrib:
                del spacing.attrib[qualified]


def _begin_field(paragraph, instruction_text: str) -> None:
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = instruction_text
    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")
    run = paragraph.add_run()
    run._r.extend([field_begin, instruction, field_separator])


def _end_field(paragraph) -> None:
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(field_end)


def _append_field(paragraph, instruction_text: str, placeholder: str = "") -> None:
    _begin_field(paragraph, instruction_text)
    if placeholder:
        placeholder_run = paragraph.add_run(placeholder)
        _font_name(placeholder_run)
        placeholder_run.font.size = Pt(12)
    _end_field(paragraph)


def _set_page_numbering(section, numbering_format: str, start: int = 1) -> None:
    section_properties = section._sectPr
    page_numbers = section_properties.find(qn("w:pgNumType"))
    if page_numbers is None:
        page_numbers = OxmlElement("w:pgNumType")
        section_properties.append(page_numbers)
    page_numbers.set(qn("w:fmt"), numbering_format)
    page_numbers.set(qn("w:start"), str(start))


def _set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    should_prompt = os.getenv("DOCX_UPDATE_FIELDS_ON_OPEN", "false").lower() in {
        "1",
        "true",
        "yes",
    }
    update_fields.set(qn("w:val"), "true" if should_prompt else "false")


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _configure_section(
    section,
    page_number_format: str | None = None,
    start: int = 1,
    vertical_alignment: str | None = None,
) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.59)
    section.footer_distance = Inches(0.6)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    valign = section._sectPr.find(qn("w:vAlign"))
    if vertical_alignment:
        if valign is None:
            valign = OxmlElement("w:vAlign")
            section._sectPr.append(valign)
        valign.set(qn("w:val"), vertical_alignment)
    elif valign is not None:
        section._sectPr.remove(valign)
    _clear_paragraph(section.header.paragraphs[0])
    footer = section.footer.paragraphs[0]
    _clear_paragraph(footer)
    if page_number_format is None:
        return
    _set_page_numbering(section, page_number_format, start=start)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_number_format == "decimal":
        run = footer.add_run("Page ")
        _font_name(run)
        run.font.size = Pt(10)
    _append_field(footer, "PAGE", "1" if page_number_format == "decimal" else "I")
    for run in footer.runs:
        _font_name(run)
        run.font.size = Pt(10)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _clear_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "bottom", "insideH", "left", "right", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _set_row_rule_borders(cell, first_row: bool, last_row: bool) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    borders = cell_properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        cell_properties.append(borders)
    for edge in ("top", "bottom", "start", "end", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        draw_rule = (edge == "top" and first_row) or (
            edge == "bottom" and (first_row or last_row)
        )
        border.set(qn("w:val"), "single" if draw_rule else "nil")
        if edge == "top" and first_row:
            size = "12"
        elif edge == "bottom" and last_row:
            size = "12"
        elif edge == "bottom" and first_row:
            size = "4"
        else:
            size = "0"
        border.set(qn("w:sz"), size if draw_rule else "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = Inches(widths[min(index, len(widths) - 1)])
            cell.width = width
            _set_cell_margins(cell)


def _column_widths(column_count: int) -> list[float]:
    if column_count <= 1:
        return [TABLE_WIDTH_IN]
    if column_count == 2:
        return [2.5, 3.25]
    if column_count == 3:
        return [2.05, 1.85, 1.85]
    if column_count == 4:
        return [1.7, 1.35, 1.35, 1.35]
    first = 1.65
    other = (TABLE_WIDTH_IN - first) / (column_count - 1)
    return [first] + [other] * (column_count - 1)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for font_attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{font_attribute}"), REPORT_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([fonts, color, underline])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_formatted_run(
    paragraph, text: str, bold: bool = False, italic: bool = False, superscript: bool = False
) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    _font_name(run)
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript


def _add_reference_bookmark(
    paragraph, marker: str, bookmark_name: str, bookmark_id: str
) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.append(start)
    _add_formatted_run(paragraph, marker)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.append(end)


def _add_citation_cross_reference(paragraph, label: str, bookmark_name: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for font_attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{font_attribute}"), REPORT_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    vertical_alignment = OxmlElement("w:vertAlign")
    vertical_alignment.set(qn("w:val"), "superscript")
    properties.extend([fonts, color, underline, vertical_alignment])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = f"[{label}]"
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _render_inline(
    paragraph,
    node,
    bold: bool = False,
    italic: bool = False,
    citation_targets: dict[str, tuple[str, str]] | None = None,
    enable_citations: bool = True,
) -> None:
    if isinstance(node, NavigableString):
        text = re.sub(r"\s+", " ", str(node))
        if not text:
            return
        if not enable_citations or not citation_targets:
            _add_formatted_run(paragraph, text, bold=bold, italic=italic)
            return
        last_index = 0
        for match in CITATION_MARKER.finditer(text):
            _add_formatted_run(paragraph, text[last_index : match.start()], bold, italic)
            label = match.group(1)
            target = citation_targets.get(label)
            if target:
                _add_citation_cross_reference(paragraph, label, target[0])
            else:
                _add_formatted_run(paragraph, match.group(0), bold, italic)
            last_index = match.end()
        _add_formatted_run(paragraph, text[last_index:], bold, italic)
        return
    if not isinstance(node, Tag):
        return
    if node.name in {"strong", "b"}:
        for child in node.children:
            _render_inline(
                paragraph, child, True, italic, citation_targets, enable_citations
            )
    elif node.name in {"em", "i"}:
        for child in node.children:
            _render_inline(
                paragraph, child, bold, True, citation_targets, enable_citations
            )
    elif node.name == "code":
        run = paragraph.add_run(node.get_text())
        _font_name(run, "Consolas")
        run.font.size = Pt(9)
        run.bold = bold
        run.italic = italic
    elif node.name == "a":
        url = node.get("href", "")
        label = node.get_text() or url
        if url:
            _add_hyperlink(paragraph, label, url)
        else:
            run = paragraph.add_run(label)
            _font_name(run)
    elif node.name == "br":
        paragraph.add_run().add_break()
    elif node.name != "img":
        for child in node.children:
            _render_inline(
                paragraph, child, bold, italic, citation_targets, enable_citations
            )


def _style_document(document: Document) -> None:
    _configure_section(document.sections[0], vertical_alignment="center")

    normal = document.styles["Normal"]
    _font_name(normal)
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = None
    normal.paragraph_format.space_after = None
    normal.paragraph_format.line_spacing = 1.5
    _set_auto_spacing(normal, True)

    for style_name, size in [
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        style = document.styles[style_name]
        _font_name(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = None
        style.paragraph_format.space_after = None
        style.paragraph_format.keep_with_next = True
        _set_auto_spacing(style, True)

    for list_style_name in ("List Paragraph", "List Bullet", "List Number"):
        style = document.styles[list_style_name]
        _font_name(style)
        style.font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_before = None
        style.paragraph_format.space_after = None
        style.paragraph_format.line_spacing = 1.5
        _set_auto_spacing(style, True)

    _set_update_fields(document)


def _add_cover_page(document: Document, report_title: str) -> None:
    report_type = document.add_paragraph()
    report_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_type.paragraph_format.space_after = Pt(28)
    run = report_type.add_run("CORPORATE VALUATION RESEARCH REPORT")
    _font_name(run)
    run.font.bold = True
    run.font.size = Pt(15)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(20)
    title_run = title_paragraph.add_run(report_title)
    _font_name(title_run)
    title_run.font.bold = True
    title_run.font.size = Pt(25)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(50)
    subtitle_run = subtitle.add_run(
        "Industry, Financial and Discounted Cash Flow Analysis"
    )
    _font_name(subtitle_run)
    subtitle_run.font.size = Pt(14)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.line_spacing = 1.5
    metadata.paragraph_format.space_after = Pt(48)
    metadata_run = metadata.add_run(
        "Prepared for analytical review\n"
        f"Report date: {date.today().strftime('%d %B %Y')}"
    )
    _font_name(metadata_run)
    metadata_run.font.size = Pt(12)

    limitation = document.add_paragraph()
    limitation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    limitation.paragraph_format.space_before = Pt(0)
    limitation_run = limitation.add_run(
        "Prepared from supplied data, model outputs and cited public evidence. "
        "Read with the report disclaimer."
    )
    _font_name(limitation_run)
    limitation_run.font.italic = True
    limitation_run.font.size = Pt(10)


def _add_contents_page(document: Document, headings: list[tuple[int, str]]) -> None:
    contents_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(contents_section, "upperRoman", start=1)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(38)
    heading.paragraph_format.space_after = Pt(30)
    run = heading.add_run("Contents")
    _font_name(run)
    run.font.bold = True
    run.font.size = Pt(18)

    toc_start = document.add_paragraph()
    toc_start.paragraph_format.space_after = Pt(0)
    _set_auto_spacing(toc_start, False)
    _begin_field(toc_start, 'TOC \\o "1-3" \\h \\z \\u')
    final_paragraph = toc_start
    for level, text in headings:
        entry = document.add_paragraph()
        entry.paragraph_format.left_indent = Inches({1: 0, 2: 0.28, 3: 0.56}[level])
        entry.paragraph_format.line_spacing = 1.15
        _set_auto_spacing(entry, True)
        run = entry.add_run(text)
        _font_name(run)
        run.font.size = Pt(12 if level == 1 else 11)
        if level == 1:
            run.bold = True
        final_paragraph = entry
    _end_field(final_paragraph)


def _add_body_section(document: Document) -> None:
    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(body_section, "decimal", start=1)


def _add_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)


def _optimise_image_file(source_path: Path) -> Path | None:
    max_width = _positive_int_from_env(
        "DOCX_IMAGE_MAX_WIDTH_PX", DEFAULT_IMAGE_MAX_WIDTH_PX
    )
    max_height = _positive_int_from_env(
        "DOCX_IMAGE_MAX_HEIGHT_PX", DEFAULT_IMAGE_MAX_HEIGHT_PX
    )
    try:
        with Image.open(source_path) as image:
            needs_resize = image.width > max_width or image.height > max_height
            if image.format == "PNG" and not needs_resize:
                return source_path

            image.load()
            image.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            if image.mode not in {"RGB", "RGBA", "L", "LA", "P"}:
                image = image.convert("RGBA" if "A" in image.mode else "RGB")

            output = tempfile.NamedTemporaryFile(
                prefix="docx_chart_", suffix=".png", delete=False
            )
            output_path = Path(output.name)
            output.close()
            image.save(output_path, format="PNG", optimize=True, compress_level=9)
            source_path.unlink(missing_ok=True)
            return output_path
    except (UnidentifiedImageError, OSError):
        source_path.unlink(missing_ok=True)
        return None


def _download_image(url: str) -> Path | None:
    parsed = urlparse(url)
    if parsed.scheme != "https" or (parsed.hostname or "").lower() not in _safe_hosts():
        return None
    source_file = tempfile.NamedTemporaryFile(
        prefix="docx_chart_source_", suffix=".img", delete=False
    )
    source_path = Path(source_file.name)
    source_file.close()
    try:
        response = requests.get(
            url,
            timeout=(4, 15),
            stream=True,
            headers={"User-Agent": "Coze-DOCX-Report-Renderer/1.0"},
        )
        response.raise_for_status()
        if not response.headers.get("Content-Type", "").lower().startswith("image/"):
            source_path.unlink(missing_ok=True)
            return None
        maximum_bytes = _positive_int_from_env(
            "DOCX_IMAGE_DOWNLOAD_MAX_BYTES", DEFAULT_IMAGE_DOWNLOAD_MAX_BYTES
        )
        downloaded = 0
        with source_path.open("wb") as destination:
            for chunk in response.iter_content(chunk_size=64 * 1024):
                if not chunk:
                    continue
                downloaded += len(chunk)
                if downloaded > maximum_bytes:
                    source_path.unlink(missing_ok=True)
                    return None
                destination.write(chunk)
        return _optimise_image_file(source_path)
    except requests.RequestException:
        source_path.unlink(missing_ok=True)
        return None


def _remove_temp_image(path: Path | None) -> None:
    if path is None:
        return
    try:
        temp_dir = Path(tempfile.gettempdir()).resolve()
        resolved = path.resolve()
        if temp_dir in resolved.parents or resolved.parent == temp_dir:
            resolved.unlink(missing_ok=True)
    except OSError:
        pass


def _render_image(
    document: Document, image_tag: Tag, include_images: bool, result: ConversionResult
) -> None:
    caption = image_tag.get("alt", "").strip() or "Chart"
    source = image_tag.get("src", "")
    image_path = _download_image(source) if include_images else None
    if image_path is None:
        result.skipped_image_count += 1
        note = document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.space_after = Pt(5)
        run = note.add_run(f"[Figure unavailable: {caption}]")
        _font_name(run)
        run.italic = True
        run.font.size = Pt(10)
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(TABLE_WIDTH_IN))
        result.image_count += 1
    finally:
        _remove_temp_image(image_path)
        gc.collect()


def _render_table(
    document: Document,
    source_table: Tag,
    result: ConversionResult,
    citation_targets: dict[str, tuple[str, str]],
) -> None:
    source_rows = source_table.find_all("tr")
    if not source_rows:
        return
    column_count = max(len(row.find_all(["th", "td"])) for row in source_rows)
    table = document.add_table(rows=len(source_rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = _column_widths(column_count)
    _set_table_geometry(table, widths)
    _clear_table_borders(table)

    for row_index, source_row in enumerate(source_rows):
        source_cells = source_row.find_all(["th", "td"])
        for column_index, source_cell in enumerate(source_cells):
            target = table.cell(row_index, column_index)
            _set_row_rule_borders(
                target,
                first_row=(row_index == 0),
                last_row=(row_index == len(source_rows) - 1),
            )
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = target.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index == 0 and row_index > 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            _set_auto_spacing(paragraph, False)
            for child in source_cell.children:
                _render_inline(
                    paragraph,
                    child,
                    bold=(row_index == 0),
                    citation_targets=citation_targets,
                )
            for run in paragraph.runs:
                _font_name(run)
                run.font.size = Pt(12)
                if row_index == 0:
                    run.bold = True

    result.table_count += 1


def _render_list(
    document: Document,
    list_tag: Tag,
    result: ConversionResult,
    citation_targets: dict[str, tuple[str, str]],
    level: int = 0,
) -> None:
    ordered = list_tag.name == "ol"
    start_number = int(list_tag.get("start", 1)) if ordered else 0
    style = "List Paragraph" if ordered else "List Bullet"
    for index, item in enumerate(list_tag.find_all("li", recursive=False), start=start_number):
        paragraph = document.add_paragraph(style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if ordered:
            paragraph.paragraph_format.left_indent = Inches(0.32 + (0.25 * level))
            paragraph.paragraph_format.first_line_indent = Inches(-0.32)
            number = paragraph.add_run(f"{index}. ")
            _font_name(number)
        elif level:
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
        for child in item.children:
            if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                continue
            _render_inline(paragraph, child, citation_targets=citation_targets)
        for child_list in item.find_all(["ul", "ol"], recursive=False):
            _render_list(document, child_list, result, citation_targets, level + 1)


def _reference_heading(text: str) -> bool:
    normalized = re.sub(r"^\d+[.)]?\s*", "", text.strip().lower())
    return normalized in {"references", "reference list", "sources", "references / sources"}


def _reference_targets(body: BeautifulSoup) -> tuple[dict[str, tuple[str, str]], set[int]]:
    targets: dict[str, tuple[str, str]] = {}
    reference_node_ids: set[int] = set()
    in_references = False
    next_id = 1
    for node in body.contents:
        if not isinstance(node, Tag):
            continue
        if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            if _reference_heading(node.get_text(" ", strip=True)):
                in_references = True
            elif in_references:
                in_references = False
            continue
        if in_references and node.name == "p":
            match = re.match(r"^\[(\d+)\]\s+", node.get_text(" ", strip=True))
            if match:
                label = match.group(1)
                reference_node_ids.add(id(node))
                if label not in targets:
                    targets[label] = (f"reference_{label}", str(next_id))
                    next_id += 1
    return targets, reference_node_ids


def _render_reference_paragraph(
    paragraph, node: Tag, citation_targets: dict[str, tuple[str, str]]
) -> None:
    label_match = re.match(r"^\[(\d+)\]\s+", node.get_text(" ", strip=True))
    if not label_match:
        for child in node.children:
            _render_inline(paragraph, child, enable_citations=False)
        return
    label = label_match.group(1)
    target = citation_targets.get(label)
    marker_consumed = False
    for child in node.children:
        if not marker_consumed and isinstance(child, NavigableString):
            raw_text = str(child)
            marker = re.match(r"^\s*\[(\d+)\]\s*", raw_text)
            if marker:
                if target:
                    _add_reference_bookmark(
                        paragraph, f"[{label}]", target[0], target[1]
                    )
                else:
                    _add_formatted_run(paragraph, f"[{label}]")
                remaining = raw_text[marker.end() :]
                if remaining:
                    _render_inline(
                        paragraph, NavigableString(" " + remaining), enable_citations=False
                    )
                marker_consumed = True
                continue
        _render_inline(paragraph, child, enable_citations=False)


def _heading_level(node: Tag, base_heading_level: int) -> int:
    text = node.get_text(" ", strip=True)
    if node.name == "h1":
        return 1
    if re.match(r"^\d+\.\d+(?:\.\d+)?\s", text):
        return 3
    if re.match(r"^\d+\.\s", text):
        return 2
    if text.lower() in {"bottom line", "references", "references / sources", "disclaimer"}:
        return 2
    return min(max(int(node.name[1]) - base_heading_level + 1, 1), 3)


def _render_block(
    document: Document,
    node: Tag,
    include_images: bool,
    result: ConversionResult,
    base_heading_level: int,
    citation_targets: dict[str, tuple[str, str]],
    reference_node_ids: set[int],
    state: RenderState,
) -> None:
    if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = _heading_level(node, base_heading_level)
        text = node.get_text(" ", strip=True)
        heading = document.add_heading(text, level=level)
        if level == 1:
            if state.report_count:
                heading.paragraph_format.page_break_before = True
            state.report_count += 1
            state.waiting_for_first_major_section = True
        elif level == 2:
            if state.waiting_for_first_major_section:
                state.waiting_for_first_major_section = False
            else:
                heading.paragraph_format.page_break_before = True
        for run in heading.runs:
            _font_name(run)
            run.font.size = Pt({1: 16, 2: 14, 3: 12}[level])
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        result.heading_count += 1
    elif node.name == "p":
        if id(node) in reference_node_ids:
            paragraph = document.add_paragraph()
            _render_reference_paragraph(paragraph, node, citation_targets)
        elif node.find("img"):
            paragraph = None
            for child in node.children:
                if isinstance(child, Tag) and child.name == "img":
                    _render_image(document, child, include_images, result)
                    paragraph = None
                    continue
                child_text = (
                    child.get_text(" ", strip=True)
                    if isinstance(child, Tag)
                    else str(child).strip()
                )
                if not child_text:
                    continue
                if paragraph is None:
                    paragraph = document.add_paragraph()
                _render_inline(paragraph, child, citation_targets=citation_targets)
        elif node.get_text(" ", strip=True):
            paragraph = document.add_paragraph()
            for child in node.children:
                _render_inline(paragraph, child, citation_targets=citation_targets)
    elif node.name == "table":
        _render_table(document, node, result, citation_targets)
    elif node.name in {"ul", "ol"}:
        _render_list(document, node, result, citation_targets)
    elif node.name == "hr":
        return
    elif node.name == "blockquote":
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        run = paragraph.add_run(node.get_text(" ", strip=True))
        _font_name(run)
        run.italic = True
    elif node.name == "pre":
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        run = paragraph.add_run(node.get_text())
        _font_name(run, "Consolas")
        run.font.size = Pt(9)


def convert_markdown_to_docx(
    markdown_text: str,
    title: str,
    output_path: Path,
    include_images: bool = True,
) -> ConversionResult:
    """Create a formal DOCX report; Markdown tables become editable Word tables."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _style_document(document)
    document.core_properties.subject = "Generated from Coze report Markdown"

    html = markdown.markdown(
        _normalise_report_markdown(markdown_text),
        extensions=["tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )
    body = BeautifulSoup(html, "html.parser")
    first_heading = body.find("h1")
    if title and title.strip():
        report_title = title.strip()
    elif first_heading is not None:
        report_title = first_heading.get_text(" ", strip=True)
    else:
        report_title = "Corporate Valuation Report"
    document.core_properties.title = report_title

    headings = body.find_all(re.compile(r"^h[1-6]$"))
    base_heading_level = (
        min(int(heading.name[1]) for heading in headings) if headings else 1
    )
    contents_headings = [
        (_heading_level(heading, base_heading_level), heading.get_text(" ", strip=True))
        for heading in headings
    ]
    citation_targets, reference_node_ids = _reference_targets(body)
    state = RenderState()

    _add_cover_page(document, report_title)
    _add_contents_page(document, contents_headings)
    _add_body_section(document)

    result = ConversionResult(output_path=output_path)
    for block in body.contents:
        if isinstance(block, Tag):
            _render_block(
                document,
                block,
                include_images,
                result,
                base_heading_level,
                citation_targets,
                reference_node_ids,
                state,
            )

    document.save(output_path)
    return result


def safe_file_stem(title: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 _-]+", "", title).strip().replace(" ", "_")
    cleaned = re.sub(r"_+", "_", cleaned)
    return cleaned[:80] or "generated_report"
